from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path(r"C:\Users\kaush\OneDrive\Desktop\ecom\generated_docs")
OUT.mkdir(exist_ok=True)
DOCX = OUT / "AKSHAY_2309301055_CartSphere_Project_Report.docx"


PROJECT_TITLE = "CartSphere E-Commerce Application"
PROJECT_SUBTITLE = "A Full-Stack Online Shopping and Order Management Web Application"
STUDENT_NAME = "Akshay Kaushik"
ROLL_NO = "2309301055"
SUPERVISOR = "Ms. Neetu Pundir"
DEPARTMENT = "SCHOOL OF COMPUTER SCIENCE & ENGINEERING"
UNIVERSITY = "GEETA UNIVERSITY, PANIPAT"
MONTH = "MAY, 2026"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D8DEE9"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, (16, 32, 58)),
        ("Heading 1", 16, (36, 62, 168)),
        ("Heading 2", 13, (16, 32, 58)),
        ("Heading 3", 11, (16, 32, 58)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "CartSphere Project Report"
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.runs[0].font.name = "Arial"
        header_p.runs[0].font.size = Pt(9)
        header_p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

        footer_p = section.footer.paragraphs[0]
        footer_p.add_run("Page ")
        add_page_number(footer_p)
        for r in footer_p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(100, 116, 139)


def add_centered(doc, text, size=12, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def add_para(doc, text, style=None, justify=True):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Field", True)
    set_cell_text(hdr[1], "Description", True)
    set_cell_shading(hdr[0], "EEF4FF")
    set_cell_shading(hdr[1], "EEF4FF")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    style_doc(doc)

    # Cover page intentionally has no visible running header.
    add_centered(doc, PROJECT_TITLE, 22, True, 18)
    add_centered(doc, "A PROJECT REPORT", 14, True, 20)
    add_centered(doc, "Submitted in partial fulfillment of the requirements for the award of the degree", 11, False, 2)
    add_centered(doc, "of", 11, False, 2)
    add_centered(doc, "BACHELOR OF COMPUTER APPLICATIONS", 14, True, 2)
    add_centered(doc, "in", 11, False, 2)
    add_centered(doc, "SOFTWARE ENGINEERING", 13, True, 24)
    add_centered(doc, "Submitted By:", 11, True, 2)
    add_centered(doc, f"{STUDENT_NAME} ({ROLL_NO})", 12, True, 14)
    add_centered(doc, "Under the Supervision of", 11, False, 2)
    add_centered(doc, SUPERVISOR, 12, True, 2)
    add_centered(doc, "(IT-Trainer) CodeQuotient", 11, False, 34)
    add_centered(doc, DEPARTMENT, 12, True, 2)
    add_centered(doc, UNIVERSITY, 12, True, 2)
    add_centered(doc, MONTH, 12, True, 2)
    add_page_break(doc)
    add_header_footer(doc)

    doc.add_heading("CANDIDATE'S DECLARATION", 1)
    add_para(doc, f'I hereby certify that the work which is being presented in the project report entitled "{PROJECT_SUBTITLE}" in partial fulfillment of the requirements for the award of the degree of Bachelor of Computer Applications with specialization in Software Engineering is an authentic record of my own work carried out during the academic session under the guidance of {SUPERVISOR}.')
    add_para(doc, "The matter presented in this project has not been submitted by me for the award of any other degree or diploma of this or any other institute.")
    doc.add_paragraph()
    add_para(doc, f"({STUDENT_NAME})", justify=False)
    add_para(doc, "This is to certify that the above statement made by the candidate is correct to the best of my knowledge.", justify=False)
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_text(table.cell(0, 0), SUPERVISOR, True)
    set_cell_text(table.cell(0, 1), "Prof. (Dr.) Amit Jain", True)
    set_cell_text(table.cell(1, 0), "IT-Trainer, CodeQuotient")
    set_cell_text(table.cell(1, 1), "Associate Dean, School of CSE")
    add_page_break(doc)

    doc.add_heading("ACKNOWLEDGEMENTS", 1)
    add_para(doc, "I would like to express my sincere gratitude to my project supervisor for valuable guidance, continuous support, and constructive feedback throughout the development of this project. The suggestions received during planning, implementation, testing, and documentation helped improve both the technical quality and presentation of the work.")
    add_para(doc, "I am thankful to the School of Computer Science and Engineering, Geeta University, Panipat, for providing the academic environment, infrastructure, and learning resources required to complete this full-stack web application project.")
    add_para(doc, "I also thank my faculty members, friends, and family for their encouragement and support during the development and documentation process.")
    add_page_break(doc)

    doc.add_heading("Abstract", 1)
    add_para(doc, "This report presents CartSphere, a full-stack e-commerce web application developed to provide a complete online shopping experience with product discovery, authentication, wishlist management, cart operations, order placement, invoice generation, coupon handling, profile management, and an admin dashboard for store operations.")
    add_para(doc, "The frontend of the application is developed using React.js, Vite, Tailwind CSS, Axios, React Router, and reusable UI components. The backend is implemented using Node.js and Express.js, while MongoDB with Mongoose is used for persistent database storage. The system uses JWT-based authentication, bcryptjs password hashing, Multer-based image upload, and structured REST APIs.")
    add_para(doc, "The application includes a responsive customer storefront, an administrative console, order status management, sales reporting, profile image uploads, product image uploads, coupon validation, cart checkout, and live order status refresh for users. The project demonstrates full-stack development practices including API design, database modeling, form validation, authentication, file handling, state management, responsive UI design, and deployment-oriented configuration.")
    add_page_break(doc)

    doc.add_heading("TABLE OF CONTENTS", 1)
    toc = [
        ("Chapter 1", "Introduction"),
        ("Chapter 2", "Review of Past Work and Problem Formulation"),
        ("Chapter 3", "System Design and Architecture"),
        ("Chapter 4", "Implementation"),
        ("Chapter 5", "Testing and Validation"),
        ("Chapter 6", "Conclusion and Future Scope"),
        ("References", "References"),
    ]
    add_kv_table(doc, toc)
    doc.add_heading("LIST OF FIGURES", 1)
    add_kv_table(doc, [
        ("Figure 1", "Three-tier architecture of CartSphere"),
        ("Figure 2", "Customer order workflow"),
        ("Figure 3", "Admin product and order management flow"),
    ])
    doc.add_heading("LIST OF TABLES", 1)
    add_kv_table(doc, [
        ("Table 1", "Technology stack"),
        ("Table 2", "Database collections"),
        ("Table 3", "Functional modules"),
        ("Table 4", "Testing summary"),
    ])
    add_page_break(doc)

    doc.add_heading("Chapter 1 - Introduction", 1)
    doc.add_heading("Background and Context", 2)
    add_para(doc, "The growth of digital commerce has changed how customers search for products, compare alternatives, place orders, and track deliveries. Modern e-commerce platforms are expected to provide smooth browsing, reliable cart management, secure authentication, convenient checkout, responsive interfaces, and efficient administrative tools.")
    add_para(doc, "CartSphere is designed as a complete academic full-stack e-commerce application. It demonstrates how a customer-facing online store and an admin-facing operations portal can work together through a REST API and MongoDB database.")
    doc.add_heading("Motivation", 2)
    add_para(doc, "The motivation behind CartSphere is to build a practical project that combines frontend design, backend development, database management, authentication, file uploads, validation, and deployment configuration in one integrated system.")
    doc.add_heading("Objectives of the Project", 2)
    add_bullets(doc, [
        "To design and develop a complete e-commerce web application using React.js, Node.js, Express.js, and MongoDB.",
        "To implement secure user registration, login, logout, profile management, and password recovery.",
        "To provide product browsing, searching, filtering, sorting, wishlist, cart, coupon, and checkout features.",
        "To build an admin dashboard for product management, user management, cart review, coupon management, order status control, invoice download, and sales reporting.",
        "To implement image upload for profile and product images using Multer while retaining URL-based image entry where required.",
        "To ensure responsive design across desktop and mobile devices.",
    ])
    doc.add_heading("Scope of the Project", 2)
    add_para(doc, "The scope of CartSphere includes customer account management, product catalog browsing, cart and order management, coupon validation, profile updates, wishlist storage, admin inventory controls, sales reporting, and deployment-ready build configuration.")
    add_page_break(doc)

    doc.add_heading("Chapter 2 - Review of Past Work and Problem Formulation", 1)
    doc.add_heading("E-Commerce Platform Requirements", 2)
    add_para(doc, "Popular e-commerce platforms such as Amazon, Flipkart, and Myntra highlight the importance of search, categorization, cart reliability, fast checkout, order visibility, and responsive interfaces. CartSphere implements these concepts at an academic project scale.")
    doc.add_heading("Problem Statement", 2)
    add_para(doc, "Small businesses and academic learners often require a manageable e-commerce system that demonstrates real-world workflows without the complexity of enterprise platforms. The problem addressed by CartSphere is the development of an end-to-end shopping platform that includes both customer and administrative functionality.")
    doc.add_heading("Functional Modules", 2)
    add_kv_table(doc, [
        ("Authentication", "Registration, login, logout, session awareness, password recovery, and protected routes."),
        ("Catalog", "Product listing, search, category filters, sorting, product detail pages, and product reviews."),
        ("Cart and Checkout", "Add to cart, quantity update, item removal, coupon validation, delivery form, payment method selection, and order placement."),
        ("Orders", "Order list, invoice download, cancellation, return request, and live status refresh."),
        ("Admin", "Dashboard statistics, product CRUD, image upload, user management, cart inspection, coupon CRUD, sales report, and order status update."),
    ])
    add_page_break(doc)

    doc.add_heading("Chapter 3 - System Design and Architecture", 1)
    doc.add_heading("System Architecture", 2)
    add_para(doc, "CartSphere follows a three-tier architecture consisting of a React presentation tier, an Express application tier, and a MongoDB data tier. The frontend communicates with the backend through REST API requests using Axios. The backend performs validation, authentication, database operations, and file upload handling.")
    add_kv_table(doc, [
        ("Presentation Tier", "React.js SPA built with Vite, Tailwind CSS, React Router, Axios, and reusable components."),
        ("Application Tier", "Node.js and Express.js backend containing routes, controllers, middleware, JWT authentication, Multer upload handling, and business logic."),
        ("Data Tier", "MongoDB database managed through Mongoose models for users, products, carts, orders, coupons, and wishlists."),
    ])
    doc.add_heading("Database Design", 2)
    add_kv_table(doc, [
        ("User", "Stores name, email, password hash, role, profile image, address, wishlist, and session state."),
        ("Product", "Stores product name, price, image, description, category, brand, stock, discount, rating, reviews, tags, and active status."),
        ("Cart", "Stores user reference and product items with quantity."),
        ("Order", "Stores user reference, ordered products, shipping address, payment method, payment status, invoice number, status, and totals."),
        ("Coupon", "Stores coupon code, description, discount type, discount value, minimum order amount, and active flag."),
    ])
    doc.add_heading("Data Flow", 2)
    add_numbered(doc, [
        "Customer registers or logs in using the authentication module.",
        "Customer browses the product catalog and adds products to wishlist or cart.",
        "Cart controller validates stock and updates cart items in MongoDB.",
        "Checkout validates delivery details, applies coupon if provided, creates an order, updates stock, and clears the cart.",
        "Admin reviews orders and updates order status; user order page refreshes to show the latest status.",
    ])
    add_page_break(doc)

    doc.add_heading("Chapter 4 - Implementation", 1)
    doc.add_heading("Technology Stack", 2)
    add_kv_table(doc, [
        ("Frontend", "React.js, Vite, Tailwind CSS, React Router, Axios, Lucide Icons."),
        ("Backend", "Node.js, Express.js, JWT, bcryptjs, Multer, CORS, dotenv."),
        ("Database", "MongoDB Atlas with Mongoose ODM."),
        ("Build and Deployment", "Vite production build, Render configuration, environment variables."),
    ])
    doc.add_heading("Frontend Implementation", 2)
    add_para(doc, "The frontend is structured into pages and reusable components. Pages include Home, Product Details, Cart, Orders, Login, Register, Profile, Wishlist, and Admin Dashboard. Components include Navbar, Footer, Product Card, Cart Item, Sidebar, Pagination, and Toast Provider.")
    add_para(doc, "The UI has been designed with responsive layouts, compact mobile category chips, polished cart cards, collapsible admin panels, and consistent form validation. Product cards support wishlist and cart actions, while the cart page provides coupon application, payment selection, and delivery detail validation.")
    doc.add_heading("Backend Implementation", 2)
    add_para(doc, "The backend follows a modular route-controller-model structure. Routes define API endpoints, controllers contain business logic, models define database schemas, and middleware handles authentication and file uploads.")
    add_para(doc, "Authentication uses JWT tokens and bcryptjs password hashing. Multer handles profile and product image uploads. Product creation supports both URL-based images and file upload, where uploaded files receive priority when both are provided.")
    doc.add_heading("Admin Dashboard", 2)
    add_para(doc, "The admin dashboard provides dashboard statistics, product management, coupon management, user management, cart review, review photo inspection, sales report export, order invoice download, and order status controls. Long tables are hidden behind show/hide panels to keep the admin interface stable and clean.")
    doc.add_heading("Validation and Security", 2)
    add_bullets(doc, [
        "Email validation for login, registration, and password recovery.",
        "Password minimum length validation.",
        "Delivery form validation for phone number, postal code, address, city, and state.",
        "Backend validation for product price, stock, discount, and required description.",
        "Protected routes for authenticated and admin-only operations.",
        "Password hashing using bcryptjs.",
    ])
    add_page_break(doc)

    doc.add_heading("Chapter 5 - Testing and Validation", 1)
    doc.add_heading("Testing Approach", 2)
    add_para(doc, "Testing was performed through manual workflow verification and API smoke testing. The application was checked for authentication, product browsing, cart updates, checkout, order creation, admin order status updates, product image upload, and responsive UI behavior.")
    add_kv_table(doc, [
        ("Login validation", "Invalid email formats and missing passwords are rejected."),
        ("Registration validation", "Name, valid email, and password length are required."),
        ("Cart workflow", "Products can be added, updated, removed, and checked out after stock validation."),
        ("Order workflow", "Invalid delivery details are rejected and valid orders are saved successfully."),
        ("Admin workflow", "Admin can add products, upload images, manage coupons, view users, inspect carts, and update orders."),
        ("Responsive UI", "Home categories, cart layout, product cards, and admin panels were adjusted for mobile and desktop use."),
    ])
    doc.add_heading("Build Verification", 2)
    add_para(doc, "The frontend production build was verified using npm run build. Backend smoke tests confirmed MongoDB connectivity, product seed operations, login validation, order validation, and payment method handling.")
    add_page_break(doc)

    doc.add_heading("Chapter 6 - Conclusion and Future Scope", 1)
    doc.add_heading("Conclusion", 2)
    add_para(doc, "CartSphere successfully demonstrates the development of a full-stack e-commerce application with a customer storefront and admin operations portal. The system integrates product management, cart checkout, order processing, coupon management, user profiles, image uploads, authentication, validation, and responsive UI design into one cohesive project.")
    add_para(doc, "The project highlights the practical application of full-stack JavaScript development and provides a strong foundation for further enhancements.")
    doc.add_heading("Future Scope", 2)
    add_bullets(doc, [
        "Integration with real payment gateways such as Razorpay or Stripe.",
        "Advanced inventory alerts and supplier management.",
        "Product recommendation system based on browsing or purchase history.",
        "Email and SMS notifications for order status changes.",
        "Analytics dashboards with charts for revenue, category performance, and customer retention.",
        "Progressive Web App support for mobile installation.",
    ])
    add_page_break(doc)

    doc.add_heading("References", 1)
    refs = [
        "React.js Documentation, https://react.dev",
        "Node.js Documentation, https://nodejs.org",
        "Express.js Documentation, https://expressjs.com",
        "MongoDB Documentation, https://www.mongodb.com/docs",
        "Mongoose Documentation, https://mongoosejs.com",
        "Vite Documentation, https://vitejs.dev",
        "Tailwind CSS Documentation, https://tailwindcss.com",
    ]
    add_numbered(doc, refs)

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    print(build())
